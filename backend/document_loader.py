import os
import logging
from pathlib import Path
from typing import Dict, Optional, Any
from datetime import datetime
import io

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import fitz
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

from backend.config import settings

logger = logging.getLogger(__name__)


SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc'}


class DocumentLoader:
    def __init__(self, tesseract_cmd: Optional[str] = None):
        self.tesseract_cmd = tesseract_cmd or settings.tesseract_cmd
        self.enable_ocr = settings.enable_ocr
        
        if self.tesseract_cmd and pytesseract:
            pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd
    
    def is_supported(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in SUPPORTED_EXTENSIONS
    
    def load(self, file_path: str, folder_source: str) -> Dict[str, Any]:
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = path.suffix.lower()
        
        if ext == '.pdf':
            return self._load_pdf(path, folder_source)
        elif ext in ['.docx', '.doc']:
            return self._load_docx(path, folder_source)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _load_pdf(self, path: Path, folder_source: str) -> Dict[str, Any]:
        content_parts = []
        metadata = {}
        
        try:
            if fitz:
                doc = fitz.open(path)
                metadata['page_count'] = len(doc)
                
                for page_num, page in enumerate(doc):
                    text = page.get_text()
                    if text.strip():
                        content_parts.append(text)
                    
                    for table in page.find_tables():
                        table_text = table.extract()
                        if table_text:
                            for row in table_text:
                                row_text = ' | '.join([cell or '' for cell in row])
                                if row_text.strip():
                                    content_parts.append(row_text)
                
                doc.close()
            elif pypdf:
                reader = pypdf.PdfReader(path)
                metadata['page_count'] = len(reader.pages)
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content_parts.append(text)
        except Exception as e:
            logger.error(f"Error loading PDF {path}: {e}")
            raise
        
        content = '\n\n'.join(content_parts)
        
        return self._create_result(path, content, metadata, folder_source)
    
    def _load_docx(self, path: Path, folder_source: str) -> Dict[str, Any]:
        if not Document:
            raise ImportError("python-docx is not installed")
        
        content_parts = []
        metadata = {}
        
        try:
            doc = Document(path)
            metadata['paragraph_count'] = len(doc.paragraphs)
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content_parts.append(text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        content_parts.append(row_text)
        except Exception as e:
            logger.error(f"Error loading DOCX {path}: {e}")
            raise
        
        content = '\n\n'.join(content_parts)
        
        return self._create_result(path, content, metadata, folder_source)
    
    def _load_image(self, path: Path, folder_source: str) -> Dict[str, Any]:
        if not Image or not pytesseract:
            raise ImportError("Pillow and pytesseract are required for image OCR")
        
        if not self.enable_ocr:
            return self._create_result(path, "", {"ocr_disabled": True}, folder_source)
        
        content = ""
        metadata = {"ocr": True}
        
        try:
            img = Image.open(path)
            
            if img.mode != 'L':
                img = img.convert('L')
            
            text = pytesseract.image_to_string(img, lang='deu+eng')
            
            content = text.strip()
            
            metadata['image_width'] = img.width
            metadata['image_height'] = img.height
        except Exception as e:
            logger.error(f"Error performing OCR on {path}: {e}")
            raise
        
        return self._create_result(path, content, metadata, folder_source)
    
    def _create_result(self, path: Path, content: str, metadata: Dict, folder_source: str) -> Dict[str, Any]:
        stat = path.stat()
        
        result = {
            'content': content,
            'metadata': {
                'source_path': str(path),
                'filename': path.name,
                'file_type': path.suffix.lower(),
                'size': stat.st_size,
                'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'folder_source': folder_source,
                **metadata
            }
        }
        
        return result


def get_all_files(folder_path: str, recursive: bool = True) -> list[str]:
    path = Path(folder_path)
    
    if not path.exists():
        raise FileNotFoundError(f"Folder not found: {folder_path}")
    
    files = []
    patterns = ['*.pdf', '*.docx', '*.doc', '*.jpg', '*.jpeg', '*.png', '*.tiff', '*.tif', '*.bmp']
    
    for pattern in patterns:
        if recursive:
            files.extend(path.rglob(pattern))
        else:
            files.extend(path.glob(pattern))
    
    return [str(f) for f in files]


_loader = None


def get_loader() -> DocumentLoader:
    global _loader
    if _loader is None:
        _loader = DocumentLoader()
    return _loader